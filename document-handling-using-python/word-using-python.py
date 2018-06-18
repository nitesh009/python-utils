# pip install python-docx
# document object contains paragraph objects whole etc..
import docx

d = docx.Document('/Users/Username/Downloads/python_practice/demo.docx')

d.paragraphs

print(d.paragraphs[0].text)

print(d.paragraphs[1].text)

p = d.paragraphs[1]
print(p.runs)

print(p.runs[0].text)

print(p.runs[1].text)

print(p.runs[2].text)

print(p.runs[3].text)

print(p.runs[1].bold)

print(p.runs[3].italic)

print(p.runs[2].italic)

p.runs[3].underline = True

p.runs[3].text = 'italic and underlined.'

d.save('/Users/Username/Downloads/demo2.docx')


d = docx.Document() # creates new document

d.add_paragraph('Hello this is a paragraph')


d.save('/Users/Username/Downloads/demo3.docx')

p = d.paragraphs[0]
p.add_run('This is a new run')

p.runs

p.runs[1].bold = True

d.save('/Users/Username/Downloads/demo4.docx')